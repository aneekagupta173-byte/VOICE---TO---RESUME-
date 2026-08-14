"""
resume_builder.py
Renders the structured résumé dict into a clean, downloadable .docx file
using python-docx. Kept deliberately simple: one clear heading hierarchy,
consistent spacing, no exotic styling that breaks across Word versions.
"""

import io

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_resume_docx(resume: dict) -> bytes:
    doc = Document()

    # Base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # Name header
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(resume.get("name", "Your Name"))
    run.bold = True
    run.font.size = Pt(22)

    # Contact line
    contact_bits = [
        resume.get("email", ""), resume.get("phone", ""), resume.get("location", "")
    ]
    contact_line = " | ".join(b for b in contact_bits if b)
    if contact_line:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.add_run(contact_line).font.size = Pt(10)

    doc.add_paragraph()  # spacer

    def add_section_heading(text):
        p = doc.add_paragraph()
        r = p.add_run(text.upper())
        r.bold = True
        r.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        # simple underline rule via bottom border would need XML; keep it clean instead
        return p

    # Summary
    if resume.get("summary"):
        add_section_heading("Summary")
        doc.add_paragraph(resume["summary"])

    # Experience
    if resume.get("experience"):
        add_section_heading("Experience")
        for job in resume["experience"]:
            p = doc.add_paragraph()
            title_run = p.add_run(job.get("title", ""))
            title_run.bold = True
            company = job.get("company", "")
            duration = job.get("duration", "")
            meta = "  —  " + ", ".join(x for x in [company, duration] if x)
            if company or duration:
                p.add_run(meta).italic = True
            for bullet in job.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.add_run(bullet)

    # Education
    if resume.get("education"):
        add_section_heading("Education")
        for edu in resume["education"]:
            p = doc.add_paragraph()
            deg_run = p.add_run(edu.get("degree", ""))
            deg_run.bold = True
            meta = ", ".join(x for x in [edu.get("institution", ""), edu.get("year", "")] if x)
            if meta:
                p.add_run("  —  " + meta).italic = True

    # Skills
    if resume.get("skills"):
        add_section_heading("Skills")
        doc.add_paragraph(" • ".join(resume["skills"]))

    # margins
    for section in doc.sections:
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
